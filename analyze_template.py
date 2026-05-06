from docx import Document
import sys

try:
    doc = Document('TEMPLATE_CV_EXPERT.docx')
    
    print(f"Total tables: {len(doc.tables)}")
    print("\n" + "="*80)
    
    # Analyze each table
    for table_idx, table in enumerate(doc.tables):
        print(f"\n### TABLE {table_idx} ###")
        print(f"Rows: {len(table.rows)}, Columns: {len(table.columns) if table.rows else 0}")
        
        # Show first 15 rows
        for row_idx, row in enumerate(table.rows[:15]):
            print(f"\nRow {row_idx}:")
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text.strip()[:100]  # First 100 chars
                if text:
                    print(f"  Cell[{cell_idx}]: {text}")
        
        if table_idx >= 4:  # Only show first 5 tables
            print("\n... (more tables)")
            break
            
except Exception as e:
    print(f"Error: {e}")
    import traceback
    traceback.print_exc()
    sys.exit(1)
