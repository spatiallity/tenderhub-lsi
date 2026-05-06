"""
Generate CV with DYNAMIC project tables - adds tables as needed
Supports unlimited number of projects by duplicating project table structure
"""
import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import os
import copy

BASE_URL = "http://localhost:8000/api/v1"

def replace_in_paragraph(paragraph, replacements):
    """Replace placeholders in a paragraph"""
    for placeholder, value in replacements.items():
        if placeholder in paragraph.text:
            # Replace in runs to preserve formatting
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(value))

def copy_table_format(source_table):
    """Create a copy of table structure with same formatting"""
    # This returns the table's XML element which can be inserted
    return copy.deepcopy(source_table._element)

def fill_project_table(table, project, project_num):
    """Fill a project table with data"""
    replacements = {
        f'{{Tahun Awal Proyek {project_num}}}': project.get('waktu_mulai', '').split()[0] if project.get('waktu_mulai') else '',
        f'{{Tahun Akhir Proyek {project_num}}}': project.get('waktu_selesai', '').split()[-1] if project.get('waktu_selesai') else '',
        f'{{Nama Proyek {project_num}}}': project.get('nama_proyek', ''),
        f'{{Lokasi Proyek {project_num}}}': project.get('lokasi_proyek', ''),
        f'{{Nama Klien {project_num}}}': project.get('pengguna_jasa') or project.get('pemberi_kerja', ''),
        f'{{Nama Perusahaan Tempat Tenaga Ahli Di Hire {project_num}}}': project.get('nama_perusahaan_lain') or project.get('bersama', 'PT SUCOFINDO (Persero)'),
        f'{{Uraian deskripsi pekerjaan {project_num}}}': project.get('uraian_tugas', ''),
        f'{{Bulan Awal, Tahun {project_num}}}': project.get('waktu_mulai', ''),
        f'{{Bulan Akhir, Tahun {project_num}}}': project.get('waktu_selesai', ''),
        f'{{Posisi Penugasan {project_num}}}': project.get('posisi_penugasan') or project.get('peran', ''),
        f'{{Status Kepegawaian {project_num}}}': project.get('status_kepegawaian', 'Tidak Tetap'),
        f'{{Nomor Surat Referensi {project_num}}}': project.get('surat_referensi', '-'),
        # Generic placeholders (without numbers)
        '{Tahun Awal Proyek 1}': project.get('waktu_mulai', '').split()[0] if project.get('waktu_mulai') else '',
        '{Tahun Akhir Proyek 1}': project.get('waktu_selesai', '').split()[-1] if project.get('waktu_selesai') else '',
        '{Nama Proyek}': project.get('nama_proyek', ''),
        '{Lokasi Proyek}': project.get('lokasi_proyek', ''),
        '{Nama Klien}': project.get('pengguna_jasa') or project.get('pemberi_kerja', ''),
        '{Nama Perusahaan Tempat Tenaga Ahli Di Hire}': project.get('nama_perusahaan_lain') or project.get('bersama', 'PT SUCOFINDO (Persero)'),
        '{Uraian deskripsi pekerjaan 1}': project.get('uraian_tugas', ''),
        '{Uraian deskripsi pekerjaan 2}': '',
        '{Bulan Awal, Tahun}': project.get('waktu_mulai', ''),
        '{Bulan Akhir, Tahun}': project.get('waktu_selesai', ''),
        '{Durasi}': f"{project.get('waktu_mulai', '')} – {project.get('waktu_selesai', '')}",
        '{Posisi Penugasan}': project.get('posisi_penugasan') or project.get('peran', ''),
        '{Status Kepegawaian}': project.get('status_kepegawaian', 'Tidak Tetap'),
        '{Nomor Surat Referensi}': project.get('surat_referensi', '-'),
    }
    
    # Replace in all cells
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, replacements)

def generate_cv_from_template(expert_id, template_path, output_filename):
    """Generate CV from template for an expert with dynamic project tables"""
    
    # Check if template exists
    if not os.path.exists(template_path):
        print(f"❌ Template not found: {template_path}")
        return False
    
    # Fetch expert data
    print(f"Fetching data for Expert ID {expert_id}...")
    response = requests.get(f"{BASE_URL}/experts/{expert_id}")
    response.raise_for_status()
    expert = response.json()
    
    print(f"Generating CV for: {expert['nama']}")
    print(f"Total projects: {len(expert.get('projects', []))}")
    print(f"Loading template: {template_path}")
    
    # Load template
    doc = Document(template_path)
    
    # Prepare replacements for basic info (header table)
    replacements = {
        '{Posisi yang Diusulkan}': expert.get('posisi_diusulkan', 'Team Leader'),
        '{Nama Lengkap}': expert.get('nama', ''),
        '{Tempat Lahir}': expert.get('tempat_lahir', ''),
        '{Tanggal Lahir}': expert.get('tanggal_lahir', ''),
        '{Tempat, Tanggal Lahir}': f"{expert.get('tempat_lahir', '')}, {expert.get('tanggal_lahir', '')}",
    }
    
    # Pendidikan Formal
    pendidikan_formal = expert.get('pendidikan_formal', [])
    if pendidikan_formal:
        replacements['{Pendidikan Formal}'] = '\n'.join(pendidikan_formal)
    
    # Pendidikan Non-Formal
    pendidikan_non_formal = expert.get('pendidikan_non_formal', [])
    if pendidikan_non_formal:
        replacements['{Pendidikan Non-Formal}'] = '\n'.join(pendidikan_non_formal)
    
    # Penguasaan Bahasa
    bahasa = expert.get('penguasaan_bahasa', [])
    if bahasa:
        replacements['{Penguasaan Bahasa}'] = '\n'.join(bahasa)
    
    # Replace in header table (first table)
    if len(doc.tables) > 0:
        header_table = doc.tables[0]
        for row in header_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
    
    # Handle projects dynamically
    projects = expert.get('projects', [])
    
    if len(projects) > 0 and len(doc.tables) > 1:
        # Find the first project table (should be table index 1)
        project_table_index = 1
        first_project_table = doc.tables[project_table_index]
        
        # Fill first project table
        print(f"Filling project 1...")
        fill_project_table(first_project_table, projects[0], 1)
        
        # For additional projects, duplicate the table
        if len(projects) > 1:
            # Find where to insert new tables (after the first project table)
            # Get the table element's parent and position
            table_element = first_project_table._element
            parent = table_element.getparent()
            
            # Find the position of the first project table
            table_position = list(parent).index(table_element)
            
            # Add additional project tables
            for i in range(1, len(projects)):
                project_num = i + 1
                print(f"Adding and filling project {project_num}...")
                
                # Create a copy of the first project table
                new_table_element = copy.deepcopy(table_element)
                
                # Insert the new table after the previous one
                # Add a paragraph break before the new table
                paragraph_element = OxmlElement('w:p')
                parent.insert(table_position + (i * 2), paragraph_element)
                parent.insert(table_position + (i * 2) + 1, new_table_element)
                
                # Create a Table object from the element to work with it
                from docx.table import Table
                new_table = Table(new_table_element, doc)
                
                # Fill the new table with project data
                fill_project_table(new_table, projects[i], project_num)
    
    # Signature section (last table)
    tanggal_sekarang = datetime.now().strftime('%d %B %Y')
    signature_replacements = {
        '{Tanggal}': tanggal_sekarang,
        '{Nama}': expert.get('nama', ''),
        '{Posisi}': expert.get('posisi_diusulkan', 'Team Leader'),
    }
    
    # Replace in last table (signature)
    if len(doc.tables) > 0:
        last_table = doc.tables[-1]
        for row in last_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, signature_replacements)
    
    # Save document
    doc.save(output_filename)
    print(f"✓ CV saved to: {output_filename}")
    print(f"  - Total Projects: {len(projects)}")
    print(f"  - Project tables created: {len(projects)}")
    
    return True

if __name__ == "__main__":
    # Configuration
    expert_id = 730  # Dr. Ir. Budi Santoso
    template_path = "../TEMPLATE_CV_EXPERT.docx"
    
    print("=" * 70)
    print("CV GENERATOR - DYNAMIC PROJECT TABLES")
    print("=" * 70)
    print()
    
    try:
        # Fetch expert name first for filename
        response = requests.get(f"{BASE_URL}/experts/{expert_id}")
        expert = response.json()
        output_file = f"CV_{expert['nama'].replace(' ', '_').replace('.', '').replace(',', '')}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        success = generate_cv_from_template(expert_id, template_path, output_file)
        if success:
            print()
            print("=" * 70)
            print("✅ SUCCESS!")
            print("=" * 70)
            print(f"\n📄 File tersimpan: {output_file}")
            print("📂 Lokasi: backend/")
            print("\n✨ Tabel proyek ditambahkan secara dinamis!")
            print("   Semua proyek sudah termasuk dalam CV.")
    except Exception as e:
        print(f"\n❌ ERROR: {e}")
        import traceback
        traceback.print_exc()
