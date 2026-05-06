"""
CV Generator API
Generates CV documents from DOCX template following Sucofindo format
"""

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload
from app.core.database import get_db
from app.models.expert import Expert
from docx import Document
from io import BytesIO
from datetime import datetime
import os

router = APIRouter()

def replace_cell_text(cell, new_text, font_name='Arial', font_size=11, alignment='left'):
    """Replace entire cell text with new text and set formatting"""
    try:
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Clear all paragraphs
        for paragraph in cell.paragraphs:
            paragraph.clear()
        
        # Set new text in first paragraph
        if cell.paragraphs:
            paragraph = cell.paragraphs[0]
        else:
            # If no paragraphs, add one
            paragraph = cell.add_paragraph()
        
        # Add run with text
        run = paragraph.add_run(str(new_text))
        
        # Set font
        run.font.name = font_name
        run.font.size = Pt(font_size)
        
        # Set alignment
        if alignment == 'left':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == 'justify':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
    except Exception as e:
        print(f"Error replacing cell text: {e}")

def replace_text_in_cell(cell, old_text, new_text):
    """Replace text in a cell while preserving formatting (fallback method)"""
    try:
        # Method 1: Try to replace in runs
        replaced = False
        for paragraph in cell.paragraphs:
            # Check if old_text exists in paragraph
            if old_text in paragraph.text:
                # Replace in runs
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, str(new_text))
                        replaced = True
        
        # Method 2: If not replaced, clear and set new text
        if not replaced and old_text in cell.text:
            replace_cell_text(cell, new_text)
            
    except Exception as e:
        print(f"Error replacing text in cell: {e}")

def generate_cv_from_template(expert_data, template_path):
    """
    Generate CV from Sucofindo DOCX template
    
    Template structure:
    - Table 0: Header info (Posisi, Nama, Tempat/Tanggal Lahir, Pendidikan, dll)
    - Tables 1-N: Individual project entries with subsections a-i
    - Last Table: Signature section
    """
    
    try:
        # Load template
        doc = Document(template_path)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to load template: {str(e)}"
        )
    
    # Get data with safe defaults
    nama = expert_data.get('nama', 'Belum diisi')
    tempat_lahir = expert_data.get('tempat_lahir', 'Belum diisi')
    tanggal_lahir = expert_data.get('tanggal_lahir', 'Belum diisi')
    posisi_diusulkan = expert_data.get('posisi_diusulkan', 'Team Leader')
    
    # Format education with type checking
    pendidikan_formal = expert_data.get('pendidikan_formal', [])
    if isinstance(pendidikan_formal, list) and pendidikan_formal:
        pendidikan_formal_text = "\n".join(str(p) for p in pendidikan_formal)
    else:
        pendidikan_formal_text = "Belum diisi"
    
    # Format non-formal education with type checking
    pendidikan_non_formal = expert_data.get('pendidikan_non_formal', [])
    if isinstance(pendidikan_non_formal, list) and pendidikan_non_formal:
        pendidikan_non_formal_text = "\n".join(str(p) for p in pendidikan_non_formal)
    else:
        pendidikan_non_formal_text = "Belum diisi"
    
    # Format language skills with type checking
    penguasaan_bahasa = expert_data.get('penguasaan_bahasa', [])
    if isinstance(penguasaan_bahasa, list) and penguasaan_bahasa:
        penguasaan_bahasa_text = "\n".join(str(b) for b in penguasaan_bahasa)
    else:
        penguasaan_bahasa_text = "Bahasa Indonesia Baik\nBahasa Inggris Baik"
    
    try:
        # Replace in Table 0 (Header info)
        if len(doc.tables) > 0:
            header_table = doc.tables[0]
            print(f"[CV Generator] Processing header table with {len(header_table.rows)} rows")
            
            # Row 0: Posisi yang diusulkan
            if len(header_table.rows) > 0 and len(header_table.rows[0].cells) > 3:
                replace_cell_text(header_table.rows[0].cells[3], posisi_diusulkan,
                                font_name='Arial', font_size=11, alignment='left')
                print(f"[CV Generator] Set posisi: {posisi_diusulkan}")
            
            # Row 2: Nama Personel
            if len(header_table.rows) > 2 and len(header_table.rows[2].cells) > 3:
                replace_cell_text(header_table.rows[2].cells[3], nama,
                                font_name='Arial', font_size=11, alignment='left')
                print(f"[CV Generator] Set nama: {nama}")
            
            # Row 3: Tempat/Tanggal Lahir
            if len(header_table.rows) > 3 and len(header_table.rows[3].cells) > 3:
                replace_cell_text(header_table.rows[3].cells[3], f"{tempat_lahir}, {tanggal_lahir}",
                                font_name='Arial', font_size=11, alignment='left')
                print(f"[CV Generator] Set tempat/tanggal lahir: {tempat_lahir}, {tanggal_lahir}")
            
            # Row 4: Pendidikan Formal
            if len(header_table.rows) > 4 and len(header_table.rows[4].cells) > 3:
                replace_cell_text(header_table.rows[4].cells[3], pendidikan_formal_text,
                                font_name='Arial', font_size=11, alignment='justify')
            
            # Row 5: Pendidikan Non Formal
            if len(header_table.rows) > 5 and len(header_table.rows[5].cells) > 3:
                replace_cell_text(header_table.rows[5].cells[3], pendidikan_non_formal_text,
                                font_name='Arial', font_size=11, alignment='justify')
            
            # Row 6: Penguasaan Bahasa
            if len(header_table.rows) > 6 and len(header_table.rows[6].cells) > 3:
                replace_cell_text(header_table.rows[6].cells[3], penguasaan_bahasa_text,
                                font_name='Arial', font_size=11, alignment='left')
        
        # Replace project data in subsequent tables (Tables 1, 2, 3, ...)
        projects = expert_data.get('projects', [])
        project_tables = doc.tables[1:-1] if len(doc.tables) > 2 else []
        
        print(f"[CV Generator] Processing {len(projects)} projects into {len(project_tables)} project tables")
        
        for idx, project in enumerate(projects):
            if idx >= len(project_tables):
                break
            
            project_table = project_tables[idx]
            print(f"[CV Generator] Processing project {idx+1}: {project.get('nama_proyek', 'N/A')}")
            
            # Determine correct cell index based on table structure
            # Tables 1-2 use cell[4], Table 3+ use cell[11]
            data_cell_idx = 4 if len(project_table.columns) <= 5 else 11
            
            # Row 1: a. Nama Proyek
            if len(project_table.rows) > 1 and len(project_table.rows[1].cells) > data_cell_idx:
                replace_cell_text(project_table.rows[1].cells[data_cell_idx], 
                                project.get('nama_proyek', 'Belum diisi'),
                                font_name='Arial', font_size=11, alignment='justify')
            
            # Row 2: b. Lokasi Proyek
            if len(project_table.rows) > 2 and len(project_table.rows[2].cells) > data_cell_idx:
                replace_cell_text(project_table.rows[2].cells[data_cell_idx], 
                                project.get('lokasi_proyek', 'Belum diisi'),
                                font_name='Arial', font_size=11, alignment='justify')
            
            # Row 3: c. Pengguna Jasa
            if len(project_table.rows) > 3 and len(project_table.rows[3].cells) > data_cell_idx:
                replace_cell_text(project_table.rows[3].cells[data_cell_idx], 
                                project.get('pengguna_jasa', 'Belum diisi'),
                                font_name='Arial', font_size=11, alignment='justify')
            
            # Row 4: d. Nama Perusahaan
            if len(project_table.rows) > 4 and len(project_table.rows[4].cells) > data_cell_idx:
                replace_cell_text(project_table.rows[4].cells[data_cell_idx], 
                                project.get('nama_perusahaan', 'PT SUCOFINDO (PERSERO)'),
                                font_name='Arial', font_size=11, alignment='justify')
            
            # Row 5: e. Uraian Tugas
            if len(project_table.rows) > 5 and len(project_table.rows[5].cells) > data_cell_idx:
                replace_cell_text(project_table.rows[5].cells[data_cell_idx], 
                                project.get('uraian_tugas', 'Belum diisi'),
                                font_name='Arial', font_size=11, alignment='justify')
            
            # Row 6: f. Waktu Pelaksanaan
            if len(project_table.rows) > 6 and len(project_table.rows[6].cells) > data_cell_idx:
                waktu_mulai = project.get('waktu_mulai', '')
                waktu_selesai = project.get('waktu_selesai', '')
                waktu_text = f"{waktu_mulai}-{waktu_selesai}" if waktu_mulai and waktu_selesai else "Belum diisi"
                replace_cell_text(project_table.rows[6].cells[data_cell_idx], waktu_text,
                                font_name='Arial', font_size=11, alignment='justify')
            
            # Row 7: g. Posisi Penugasan
            if len(project_table.rows) > 7 and len(project_table.rows[7].cells) > data_cell_idx:
                replace_cell_text(project_table.rows[7].cells[data_cell_idx], 
                                project.get('posisi_penugasan', 'Belum diisi'),
                                font_name='Arial', font_size=11, alignment='justify')
            
            # Row 8: h. Status Kepegawaian
            if len(project_table.rows) > 8 and len(project_table.rows[8].cells) > data_cell_idx:
                replace_cell_text(project_table.rows[8].cells[data_cell_idx], 
                                project.get('status_kepegawaian', 'Tidak Tetap'),
                                font_name='Arial', font_size=11, alignment='justify')
            
            # Row 9: i. Surat Referensi
            if len(project_table.rows) > 9 and len(project_table.rows[9].cells) > data_cell_idx:
                replace_cell_text(project_table.rows[9].cells[data_cell_idx], 
                                project.get('surat_referensi', '-'),
                                font_name='Arial', font_size=11, alignment='justify')
        
        # Update signature table (last table)
        if len(doc.tables) > 0:
            signature_table = doc.tables[-1]
            print(f"[CV Generator] Processing signature table")
            
            # Update date
            if len(signature_table.rows) > 0 and len(signature_table.rows[0].cells) > 0:
                today = datetime.now().strftime("%d %B %Y")
                replace_cell_text(signature_table.rows[0].cells[0], f"Jakarta, {today}",
                                font_name='Arial', font_size=11, alignment='left')
            
            # Update name in signature
            if len(signature_table.rows) > 3 and len(signature_table.rows[3].cells) > 0:
                replace_cell_text(signature_table.rows[3].cells[0], nama,
                                font_name='Arial', font_size=11, alignment='left')
            
            # Update position in signature
            if len(signature_table.rows) > 4 and len(signature_table.rows[4].cells) > 0:
                replace_cell_text(signature_table.rows[4].cells[0], posisi_diusulkan,
                                font_name='Arial', font_size=11, alignment='left')
    
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error processing template: {str(e)}"
        )
    
    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output

@router.get("/{expert_id}/cv")
async def generate_expert_cv(
    expert_id: int,
    db: AsyncSession = Depends(get_db)
):
    """
    Generate CV document for an expert using Sucofindo template
    Returns a downloadable DOCX file
    """
    
    try:
        # Get expert with all related data
        result = await db.execute(
            select(Expert)
            .options(selectinload(Expert.projects), selectinload(Expert.reviews))
            .where(Expert.id == expert_id)
        )
        expert = result.scalars().first()
        
        if not expert:
            raise HTTPException(status_code=404, detail="Expert not found")
        
        print(f"[CV Generator] Generating CV for expert: {expert.nama} (ID: {expert_id})")
        
        # Prepare expert data following template structure
        expert_data = {
            'nama': expert.nama,
            'posisi_diusulkan': getattr(expert, 'posisi_diusulkan', None) or 'Team Leader',
            'tempat_lahir': getattr(expert, 'tempat_lahir', None) or 'Belum diisi',
            'tanggal_lahir': getattr(expert, 'tanggal_lahir', None) or 'Belum diisi',
            'pendidikan_formal': getattr(expert, 'pendidikan_formal', None) or [],
            'pendidikan_non_formal': getattr(expert, 'pendidikan_non_formal', None) or [],
            'penguasaan_bahasa': getattr(expert, 'penguasaan_bahasa', None) or ['Bahasa Indonesia Baik', 'Bahasa Inggris Baik'],
            'projects': [
                {
                    'nama_proyek': p.nama_proyek,
                    'lokasi_proyek': getattr(p, 'lokasi_proyek', None) or 'Belum diisi',
                    'pengguna_jasa': getattr(p, 'pengguna_jasa', None) or p.pemberi_kerja or 'Belum diisi',
                    'nama_perusahaan': p.nama_perusahaan_lain or 'PT SUCOFINDO (PERSERO)',
                    'uraian_tugas': getattr(p, 'uraian_tugas', None) or 'Belum diisi',
                    'waktu_mulai': getattr(p, 'waktu_mulai', None) or (str(p.tahun) if p.tahun else ''),
                    'waktu_selesai': getattr(p, 'waktu_selesai', None) or (str(p.tahun) if p.tahun else ''),
                    'posisi_penugasan': getattr(p, 'posisi_penugasan', None) or p.peran or 'Belum diisi',
                    'status_kepegawaian': getattr(p, 'status_kepegawaian', None) or 'Tidak Tetap',
                    'surat_referensi': getattr(p, 'surat_referensi', None) or '-',
                }
                for p in expert.projects[:3]  # Max 3 projects (template has 3 project tables)
            ]
        }
        
        print(f"[CV Generator] Expert data prepared. Projects: {len(expert_data['projects'])}")
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"[CV Generator] Error preparing expert data: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error preparing expert data: {str(e)}"
        )
    
    # Get template path - try multiple locations
    # 1. Try project root (development)
    template_path = os.path.join(
        os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))),
        'TEMPLATE_CV_EXPERT.docx'
    )
    
    # 2. If not found, try one level up (deployment)
    if not os.path.exists(template_path):
        template_path = os.path.join(
            os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))),
            'TEMPLATE_CV_EXPERT.docx'
        )
    
    # 3. If still not found, try current working directory
    if not os.path.exists(template_path):
        template_path = os.path.join(os.getcwd(), 'TEMPLATE_CV_EXPERT.docx')
    
    # 4. If still not found, try backend parent directory
    if not os.path.exists(template_path):
        template_path = os.path.join(os.path.dirname(os.getcwd()), 'TEMPLATE_CV_EXPERT.docx')
    
    print(f"[CV Generator] Template path: {template_path}")
    print(f"[CV Generator] Template exists: {os.path.exists(template_path)}")
    print(f"[CV Generator] Current working directory: {os.getcwd()}")
    
    if not os.path.exists(template_path):
        # List all attempted paths for debugging
        attempted_paths = [
            os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))), 'TEMPLATE_CV_EXPERT.docx'),
            os.path.join(os.getcwd(), 'TEMPLATE_CV_EXPERT.docx'),
            os.path.join(os.path.dirname(os.getcwd()), 'TEMPLATE_CV_EXPERT.docx'),
        ]
        raise HTTPException(
            status_code=500, 
            detail=f"CV template not found. Tried paths: {', '.join(attempted_paths)}. Current dir: {os.getcwd()}"
        )
    
    # Generate CV
    try:
        print(f"[CV Generator] Starting CV generation...")
        cv_file = generate_cv_from_template(expert_data, template_path)
        print(f"[CV Generator] CV generated successfully")
    except HTTPException:
        raise
    except Exception as e:
        print(f"[CV Generator] Error generating CV: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate CV: {str(e)}"
        )
    
    # Return as downloadable file
    filename = f"CV_{expert.nama.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
    
    return StreamingResponse(
        cv_file,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        }
    )
