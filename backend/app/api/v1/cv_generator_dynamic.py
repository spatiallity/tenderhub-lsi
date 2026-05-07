"""
CV Generator API - Dynamic Version
Generates CV documents with UNLIMITED project tables
Each project gets its own table, duplicated from template
"""

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload
from app.core.database import get_db
from app.models.expert import Expert
from docx import Document
from docx.table import Table
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from io import BytesIO
from datetime import datetime
import os
import copy
import shutil
import subprocess
import tempfile

CV_FONT_NAME = 'Arial'
CV_FONT_SIZE_PT = 11


def _force_arial(run):
    """Force Arial on a run including East Asia / complex script font slots
    (Word otherwise falls back to the document default)."""
    run.font.name = CV_FONT_NAME
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(attr), CV_FONT_NAME)

router = APIRouter()

@router.options("/{expert_id}/cv")
async def cv_options(expert_id: int):
    """Handle CORS preflight for CV generation"""
    return {}


def _find_libreoffice() -> str | None:
    """Locate the LibreOffice/soffice binary or return None."""
    for cand in ("soffice", "libreoffice"):
        path = shutil.which(cand)
        if path:
            return path
    common = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/bin/libreoffice",
        "/usr/bin/soffice",
        "/usr/local/bin/soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ]
    for p in common:
        if os.path.exists(p):
            return p
    return None


def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert DOCX bytes to PDF bytes using LibreOffice headless. Raises HTTPException on failure."""
    soffice = _find_libreoffice()
    if not soffice:
        raise HTTPException(
            status_code=503,
            detail="PDF conversion unavailable: LibreOffice (soffice) not installed on server. Install LibreOffice or use DOCX export."
        )

    with tempfile.TemporaryDirectory() as tmpdir:
        in_path = os.path.join(tmpdir, "input.docx")
        with open(in_path, "wb") as f:
            f.write(docx_bytes)

        cmd = [soffice, "--headless", "--norestore", "--nologo",
               "--convert-to", "pdf", "--outdir", tmpdir, in_path]
        try:
            proc = subprocess.run(cmd, capture_output=True, timeout=120)
        except subprocess.TimeoutExpired:
            raise HTTPException(status_code=504, detail="PDF conversion timed out (>120s).")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"LibreOffice invocation failed: {e}")

        if proc.returncode != 0:
            err = (proc.stderr or proc.stdout or b"").decode("utf-8", errors="ignore")
            raise HTTPException(status_code=500, detail=f"LibreOffice conversion failed: {err[:500]}")

        out_path = os.path.join(tmpdir, "input.pdf")
        if not os.path.exists(out_path):
            raise HTTPException(status_code=500, detail="LibreOffice did not produce a PDF file.")

        with open(out_path, "rb") as f:
            return f.read()

def replace_in_paragraph(paragraph, replacements):
    """Replace placeholders in a paragraph. Longest keys first to avoid
    substring collisions like {Posisi} clobbering {Posisi Tenaga Ahli}.
    Falls back to whole-paragraph rewrite when a placeholder is split across
    multiple runs (common Word quirk — runs get re-segmented on edits).
    Forces Arial on any run that received replacement content."""
    keys = sorted(replacements.keys(), key=len, reverse=True)
    for placeholder in keys:
        value = str(replacements[placeholder])
        if placeholder not in paragraph.text:
            continue
        replaced_in_run = False
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)
                _force_arial(run)
                replaced_in_run = True
        if replaced_in_run and placeholder not in paragraph.text:
            continue
        if placeholder in paragraph.text:
            full = paragraph.text.replace(placeholder, value)
            runs = paragraph.runs
            if runs:
                runs[0].text = full
                _force_arial(runs[0])
                for r in runs[1:]:
                    r.text = ''


def _apply_bullet_format(paragraph, line):
    """If a line starts with a bullet glyph, swap it for `•\tTEXT` and apply a
    hanging indent so wrapped lines align under the bullet's text — not under
    the dot. Also tightens line spacing so the bulleted block reads cleanly."""
    if line.startswith(('• ', '•\t', '- ', '* ')):
        body = line[2:].lstrip() if line[1] in (' ', '\t') else line[1:].lstrip()
        pf = paragraph.paragraph_format
        pf.left_indent = Cm(0.6)
        pf.first_line_indent = Cm(-0.6)
        pf.space_after = Pt(2)
        return f'•\t{body}'
    return line


def set_cell_multiline(cell, text):
    """Replace cell text with multi-line content. Splits on newlines into
    separate paragraphs. Forces Arial on every new run. Bullet lines (starting
    with `• `, `- `, or `* `) get a proper hanging indent — wrapped lines
    align under text, not under the bullet glyph."""
    paras = list(cell.paragraphs)
    first = paras[0]
    for p in paras[1:]:
        p._element.getparent().remove(p._element)
    for r in list(first.runs):
        r._element.getparent().remove(r._element)
    lines = (text or '').split('\n')
    if not lines:
        return
    formatted_first = _apply_bullet_format(first, lines[0])
    run = first.add_run(formatted_first)
    _force_arial(run)
    for line in lines[1:]:
        new_para = cell.add_paragraph()
        formatted = _apply_bullet_format(new_para, line)
        run = new_para.add_run(formatted)
        _force_arial(run)


def strip_contoh_lines(table):
    """Remove paragraphs that begin with 'Contoh:' from every cell."""
    for row in table.rows:
        for cell in row.cells:
            for p in list(cell.paragraphs):
                if p.text.strip().lower().startswith('contoh:'):
                    p._element.getparent().remove(p._element)

def fill_project_table(table, project, project_num):
    """Fill a project table with data from a single project"""
    
    # Extract year from waktu_mulai or waktu_selesai
    waktu_mulai = project.get('waktu_mulai', '')
    waktu_selesai = project.get('waktu_selesai', '')
    
    # Try to extract year
    year_start = waktu_mulai.split()[-1] if waktu_mulai else ''
    year_end = waktu_selesai.split()[-1] if waktu_selesai else ''
    year_display = f"{year_start} – {year_end}" if year_start and year_end else (year_start or year_end or '')
    
    # Format uraian tugas as bullet points
    uraian_tugas_raw = project.get('uraian_tugas', '')
    if uraian_tugas_raw and '\n' in uraian_tugas_raw:
        # Already has line breaks - split and format as bullets
        tasks = [line.strip() for line in uraian_tugas_raw.split('\n') if line.strip()]
        uraian_tugas_formatted = '\n'.join(f"• {task}" if not task.startswith('•') else task for task in tasks)
    else:
        # Single line - just add bullet
        uraian_tugas_formatted = f"• {uraian_tugas_raw}" if uraian_tugas_raw else ''
    
    replacements = {
        # Year header
        f'{{Tahun Awal Proyek {project_num}}}': year_start,
        f'{{Tahun Akhir Proyek {project_num}}}': year_end,
        '{Tahun Awal Proyek 1}': year_start,
        '{Tahun Akhir Proyek 1}': year_end,
        
        # Project details
        f'{{Nama Proyek {project_num}}}': project.get('nama_proyek', ''),
        f'{{Lokasi Proyek {project_num}}}': project.get('lokasi_proyek', ''),
        f'{{Nama Klien {project_num}}}': project.get('pengguna_jasa') or project.get('pemberi_kerja', ''),
        f'{{Nama Perusahaan Tempat Tenaga Ahli Di Hire {project_num}}}': project.get('nama_perusahaan_lain') or project.get('bersama', 'PT SUCOFINDO (Persero)'),
        f'{{Uraian deskripsi pekerjaan {project_num}}}': uraian_tugas_formatted,
        f'{{Bulan Awal, Tahun {project_num}}}': waktu_mulai,
        f'{{Bulan Akhir, Tahun {project_num}}}': waktu_selesai,
        f'{{Posisi Penugasan {project_num}}}': project.get('posisi_penugasan') or project.get('peran', ''),
        f'{{Status Kepegawaian {project_num}}}': project.get('status_kepegawaian', 'Tidak Tetap'),
        f'{{Nomor Surat Referensi {project_num}}}': project.get('surat_referensi', '-'),
        
        # Generic placeholders (without numbers) - for first project
        '{Nama Proyek}': project.get('nama_proyek', ''),
        '{Lokasi Proyek}': project.get('lokasi_proyek', ''),
        '{Nama Klien}': project.get('pengguna_jasa') or project.get('pemberi_kerja', ''),
        '{Nama Perusahaan Tempat Tenaga Ahli Di Hire}': project.get('nama_perusahaan_lain') or project.get('bersama', 'PT SUCOFINDO (Persero)'),
        '{Uraian deskripsi pekerjaan 1}': uraian_tugas_formatted,
        '{Uraian deskripsi pekerjaan 2}': '',
        '{Bulan Awal, Tahun}': waktu_mulai,
        '{Bulan Akhir, Tahun}': waktu_selesai,
        # {Durasi} would otherwise duplicate the range — leave blank, post-pass strips empty parens.
        '{Durasi}': '',
        '{Posisi Penugasan}': project.get('posisi_penugasan') or project.get('peran', ''),
        '{Status Kepegawaian}': project.get('status_kepegawaian', 'Tidak Tetap'),
        '{Nomor Surat Referensi}': project.get('surat_referensi', '-'),
    }
    
    # Replace in all cells
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, replacements)

    # Strip 'Contoh: …' helper rows the template ships with.
    strip_contoh_lines(table)

    # Collapse empty parens left behind when {Durasi} resolved to ''.
    import re as _re
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if '()' in p.text or '( )' in p.text:
                    new = _re.sub(r'\s*\(\s*\)', '', p.text)
                    if new != p.text:
                        runs = p.runs
                        if runs:
                            runs[0].text = new
                            for r in runs[1:]:
                                r.text = ''

    # Inject Uraian Tugas content by row label (template ships an empty value cell).
    uraian_tugas_raw = (project.get('uraian_tugas') or '').strip()
    if uraian_tugas_raw:
        if '\n' in uraian_tugas_raw:
            lines = [l.strip() for l in uraian_tugas_raw.split('\n') if l.strip()]
            uraian_text = '\n'.join(l if l.startswith('•') else f'• {l}' for l in lines)
        else:
            uraian_text = f'• {uraian_tugas_raw}'
        for row in table.rows:
            if len(row.cells) >= 5 and row.cells[2].text.strip().lower() == 'uraian tugas':
                set_cell_multiline(row.cells[4], uraian_text)
                break

def generate_cv_from_template_dynamic(expert_data, template_path):
    """
    Generate CV from Sucofindo DOCX template with DYNAMIC project tables
    
    Template structure:
    - Table 0: Header info (Posisi, Nama, Tempat/Tanggal Lahir, Pendidikan, dll)
    - Table 1: First project (template) - will be duplicated for each project
    - Last Table: Signature section
    
    This function will:
    1. Fill header table with personal info
    2. Fill first project table
    3. Duplicate project table for each additional project
    4. Fill signature table
    """
    
    try:
        # Load template
        doc = Document(template_path)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to load template: {str(e)}"
        )
    
    # Get data with safe defaults
    nama = expert_data.get('nama', 'Belum diisi')
    tempat_lahir = expert_data.get('tempat_lahir', 'Belum diisi')
    tanggal_lahir = expert_data.get('tanggal_lahir', 'Belum diisi')
    posisi_diusulkan = expert_data.get('posisi_diusulkan', 'Team Leader')
    
    # Format education
    pendidikan_formal = expert_data.get('pendidikan_formal', [])
    if isinstance(pendidikan_formal, list) and pendidikan_formal:
        pendidikan_formal_text = "\n".join(f"• {str(p)}" for p in pendidikan_formal)
    else:
        pendidikan_formal_text = ""
    
    pendidikan_non_formal = expert_data.get('pendidikan_non_formal', [])
    if isinstance(pendidikan_non_formal, list) and pendidikan_non_formal:
        pendidikan_non_formal_text = "\n".join(f"• {str(p)}" for p in pendidikan_non_formal)
    else:
        pendidikan_non_formal_text = ""
    
    penguasaan_bahasa = expert_data.get('penguasaan_bahasa', [])
    if isinstance(penguasaan_bahasa, list) and penguasaan_bahasa:
        penguasaan_bahasa_text = "\n".join(f"• {str(b)}" for b in penguasaan_bahasa)
    else:
        penguasaan_bahasa_text = "• Bahasa Indonesia: Sangat Baik\n• Bahasa Inggris: Baik"
    
    # Prepare replacements for header table — match TEMPLATE_CV_EXPERT.docx exactly.
    header_replacements = {
        # Current template names
        '{Posisi}': posisi_diusulkan,
        '{Nama Tenaga Ahli}': nama,
        '{Tempat}': tempat_lahir,
        '{Hari, Tanggal Bulan Tahun}': tanggal_lahir,
        # Legacy placeholders (kept for older template variants)
        '{Posisi yang Diusulkan}': posisi_diusulkan,
        '{Nama Lengkap}': nama,
        '{Tempat Lahir}': tempat_lahir,
        '{Tanggal Lahir}': tanggal_lahir,
        '{Tempat, Tanggal Lahir}': f"{tempat_lahir}, {tanggal_lahir}",
        '{Pendidikan Formal}': pendidikan_formal_text,
        '{Pendidikan Non-Formal}': pendidikan_non_formal_text,
        '{Penguasaan Bahasa}': penguasaan_bahasa_text,
    }

    # Header rows whose value cell is empty in the template — fill by label.
    header_label_to_text = {
        'pendidikan': pendidikan_formal_text,
        'pendidikan non formal': pendidikan_non_formal_text,
        'pendidikan non-formal': pendidikan_non_formal_text,
        'penguasaan bahasa': penguasaan_bahasa_text,
    }
    
    try:
        # Replace in header table (Table 0)
        if len(doc.tables) > 0:
            header_table = doc.tables[0]
            print(f"[CV Generator Dynamic] Processing header table")
            for row in header_table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, header_replacements)

            # Fill rows whose value cell is empty in the template (Pendidikan, Bahasa, …)
            for row in header_table.rows:
                if len(row.cells) < 4:
                    continue
                label = row.cells[1].text.strip().lower()
                if label in header_label_to_text:
                    text = header_label_to_text[label]
                    if text and not row.cells[3].text.strip():
                        set_cell_multiline(row.cells[3], text)
        
        # Handle projects dynamically
        projects = expert_data.get('projects', [])
        print(f"[CV Generator Dynamic] Processing {len(projects)} projects")
        
        if len(projects) > 0 and len(doc.tables) > 1:
            # Find the first project table (should be table index 1)
            project_table_index = 1
            first_project_table = doc.tables[project_table_index]
            
            # Fill first project table
            print(f"[CV Generator Dynamic] Filling project 1")
            fill_project_table(first_project_table, projects[0], 1)
            
            # For additional projects, duplicate the table
            if len(projects) > 1:
                # Get the table element's parent and position
                table_element = first_project_table._element
                parent = table_element.getparent()
                
                # Find the position of the first project table
                table_position = list(parent).index(table_element)
                
                # Add additional project tables
                for i in range(1, len(projects)):
                    project_num = i + 1
                    print(f"[CV Generator Dynamic] Adding and filling project {project_num}")
                    
                    # Create a deep copy of the first project table
                    new_table_element = copy.deepcopy(table_element)
                    
                    # Insert the new table after the previous one
                    # Add a paragraph break before the new table for spacing
                    from docx.oxml import OxmlElement
                    paragraph_element = OxmlElement('w:p')
                    insert_position = table_position + (i * 2)
                    parent.insert(insert_position, paragraph_element)
                    parent.insert(insert_position + 1, new_table_element)
                    
                    # Create a Table object from the element to work with it
                    new_table = Table(new_table_element, doc)
                    
                    # Fill the new table with project data
                    fill_project_table(new_table, projects[i], project_num)
        
        # Update signature table (last table)
        if len(doc.tables) > 0:
            signature_table = doc.tables[-1]
            print(f"[CV Generator Dynamic] Processing signature table")
            
            tanggal_sekarang = datetime.now().strftime('%d %B %Y')
            # Length-sorted replace inside replace_in_paragraph protects
            # {Posisi Tenaga Ahli} from being clobbered by {Posisi}.
            signature_replacements = {
                # Current template
                '{Nama Tenaga Ahli}': nama,
                '{Posisi Tenaga Ahli}': posisi_diusulkan,
                '{Tangal Bulan Tahun}': tanggal_sekarang,   # template typo "Tangal"
                '{Tanggal Bulan Tahun}': tanggal_sekarang,  # tolerate fixed spelling
                # Legacy
                '{Tanggal}': tanggal_sekarang,
                '{Nama}': nama,
                '{Posisi}': posisi_diusulkan,
            }
            
            for row in signature_table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, signature_replacements)
    
    except Exception as e:
        print(f"[CV Generator Dynamic] Error: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Error processing template: {str(e)}"
        )
    
    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    print(f"[CV Generator Dynamic] CV generated successfully with {len(projects)} project tables")
    return output

async def _build_cv_for_expert(expert_id: int, db: AsyncSession):
    """Shared logic: load expert + render DOCX. Returns (BytesIO, expert)."""
    try:
        # Get expert with all related data
        result = await db.execute(
            select(Expert)
            .options(selectinload(Expert.projects), selectinload(Expert.reviews))
            .where(Expert.id == expert_id)
        )
        expert = result.scalars().first()
        
        if not expert:
            raise HTTPException(status_code=404, detail="Expert not found")
        
        print(f"[CV Generator Dynamic] Generating CV for expert: {expert.nama} (ID: {expert_id})")
        print(f"[CV Generator Dynamic] Total projects: {len(expert.projects)}")
        
        # Prepare expert data
        expert_data = {
            'nama': expert.nama,
            'posisi_diusulkan': getattr(expert, 'posisi_diusulkan', None) or 'Team Leader',
            'tempat_lahir': getattr(expert, 'tempat_lahir', None) or 'Belum diisi',
            'tanggal_lahir': getattr(expert, 'tanggal_lahir', None) or 'Belum diisi',
            'pendidikan_formal': getattr(expert, 'pendidikan_formal', None) or [],
            'pendidikan_non_formal': getattr(expert, 'pendidikan_non_formal', None) or [],
            'penguasaan_bahasa': getattr(expert, 'penguasaan_bahasa', None) or ['Bahasa Indonesia Baik', 'Bahasa Inggris Baik'],
            'projects': [
                {
                    'nama_proyek': p.nama_proyek,
                    'lokasi_proyek': getattr(p, 'lokasi_proyek', None) or 'Belum diisi',
                    'pengguna_jasa': getattr(p, 'pengguna_jasa', None) or p.pemberi_kerja or 'Belum diisi',
                    'nama_perusahaan_lain': p.nama_perusahaan_lain or 'PT SUCOFINDO (PERSERO)',
                    'bersama': p.bersama or 'PT SUCOFINDO (PERSERO)',
                    'uraian_tugas': getattr(p, 'uraian_tugas', None) or 'Belum diisi',
                    'waktu_mulai': getattr(p, 'waktu_mulai', None) or (str(p.tahun) if p.tahun else ''),
                    'waktu_selesai': getattr(p, 'waktu_selesai', None) or (str(p.tahun) if p.tahun else ''),
                    'posisi_penugasan': getattr(p, 'posisi_penugasan', None) or p.peran or 'Belum diisi',
                    'status_kepegawaian': getattr(p, 'status_kepegawaian', None) or 'Tidak Tetap',
                    'surat_referensi': getattr(p, 'surat_referensi', None) or '-',
                    'peran': p.peran,
                    'pemberi_kerja': p.pemberi_kerja,
                }
                for p in expert.projects  # ✅ ALL projects, not limited to 3
            ]
        }
        
        print(f"[CV Generator Dynamic] Expert data prepared. Projects: {len(expert_data['projects'])}")
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"[CV Generator Dynamic] Error preparing expert data: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error preparing expert data: {str(e)}"
        )
    
    # Get template path - try multiple locations
    # In HuggingFace: working dir is /home/user/app, file is at /home/user/app/TEMPLATE_CV_EXPERT.docx
    # In local dev: file is at backend/TEMPLATE_CV_EXPERT.docx
    
    possible_paths = [
        # HuggingFace: /home/user/app/TEMPLATE_CV_EXPERT.docx
        os.path.join(os.getcwd(), 'TEMPLATE_CV_EXPERT.docx'),
        # Relative to this file: backend/app/api/v1/ -> backend/
        os.path.join(os.path.dirname(__file__), '..', '..', '..', 'TEMPLATE_CV_EXPERT.docx'),
        os.path.join(os.path.dirname(__file__), '..', '..', 'TEMPLATE_CV_EXPERT.docx'),
        os.path.join(os.path.dirname(__file__), '..', 'TEMPLATE_CV_EXPERT.docx'),
        # Absolute fallback
        '/home/user/app/TEMPLATE_CV_EXPERT.docx',
        '/app/TEMPLATE_CV_EXPERT.docx',
    ]
    
    template_path = None
    for p in possible_paths:
        resolved = os.path.abspath(p)
        print(f"[CV Generator Dynamic] Checking template path: {resolved}")
        if os.path.exists(resolved):
            template_path = resolved
            print(f"[CV Generator Dynamic] ✅ Template found at: {resolved}")
            break
    
    if not template_path:
        # List files in cwd for debugging
        try:
            cwd_files = os.listdir(os.getcwd())
            print(f"[CV Generator Dynamic] Files in cwd ({os.getcwd()}): {cwd_files}")
        except Exception:
            pass
        raise HTTPException(
            status_code=500, 
            detail=f"CV template not found. Searched: {[os.path.abspath(p) for p in possible_paths]}"
        )
    
    # Generate CV with dynamic tables
    try:
        print(f"[CV Generator Dynamic] Starting CV generation...")
        cv_file = generate_cv_from_template_dynamic(expert_data, template_path)
        print(f"[CV Generator Dynamic] CV generated successfully")
    except HTTPException:
        raise
    except Exception as e:
        print(f"[CV Generator Dynamic] Error generating CV: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate CV: {str(e)}"
        )
    
    return cv_file, expert


@router.get("/{expert_id}/cv")
async def generate_expert_cv_dynamic(
    expert_id: int,
    db: AsyncSession = Depends(get_db)
):
    """
    Generate CV document for an expert using Sucofindo template.
    Supports UNLIMITED number of projects (dynamic table generation).
    Returns a downloadable DOCX file.
    """
    cv_file, expert = await _build_cv_for_expert(expert_id, db)
    filename = f"CV_{expert.nama.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"

    from fastapi.responses import Response
    return Response(
        content=cv_file.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Cache-Control": "no-cache",
        }
    )


@router.options("/{expert_id}/cv/pdf")
async def cv_pdf_options(expert_id: int):
    """CORS preflight for PDF route."""
    return {}


@router.get("/{expert_id}/cv/pdf")
async def generate_expert_cv_pdf(
    expert_id: int,
    db: AsyncSession = Depends(get_db)
):
    """
    Generate expert CV as PDF (DOCX rendered then converted via LibreOffice headless).
    Returns 503 if LibreOffice is not installed on the server.
    """
    cv_file, expert = await _build_cv_for_expert(expert_id, db)
    print(f"[CV Generator Dynamic] Converting to PDF via LibreOffice…")
    pdf_bytes = convert_docx_to_pdf(cv_file.getvalue())
    filename = f"CV_{expert.nama.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf"

    from fastapi.responses import Response
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Cache-Control": "no-cache",
        }
    )
