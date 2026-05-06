"""
Generate CV from TEMPLATE_CV_EXPERT.docx by filling in the placeholders
"""
import requests
from docx import Document
from datetime import datetime
import os

BASE_URL = "http://localhost:8000/api/v1"

def format_currency(value):
    """Format currency to Indonesian Rupiah"""
    if not value:
        return "Rp 0"
    return f"Rp {value:,.0f}".replace(",", ".")

def replace_placeholder(doc, placeholder, value):
    """Replace placeholder in all paragraphs and tables"""
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(value))
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, str(value))

def add_table_row(table, data_list):
    """Add a new row to table with data"""
    row = table.add_row()
    for i, data in enumerate(data_list):
        if i < len(row.cells):
            row.cells[i].text = str(data)

def generate_cv_from_template(expert_id, template_path, output_filename):
    """Generate CV from template for an expert"""
    
    # Check if template exists
    if not os.path.exists(template_path):
        print(f"❌ Template not found: {template_path}")
        return False
    
    # Fetch expert data
    print(f"Fetching data for Expert ID {expert_id}...")
    response = requests.get(f"{BASE_URL}/experts/{expert_id}")
    response.raise_for_status()
    expert = response.json()
    
    print(f"Generating CV for: {expert['nama']}")
    print(f"Loading template: {template_path}")
    
    # Load template
    doc = Document(template_path)
    
    # Replace basic placeholders
    replacements = {
        '{{NAMA}}': expert.get('nama', ''),
        '{{TEMPAT_LAHIR}}': expert.get('tempat_lahir', ''),
        '{{TANGGAL_LAHIR}}': expert.get('tanggal_lahir', ''),
        '{{NO_HP}}': expert.get('no_hp', ''),
        '{{INSTANSI}}': expert.get('instansi', ''),
        '{{KEAHLIAN}}': ', '.join(expert.get('keahlian', [])),
        '{{PORTFOLIO}}': ', '.join(expert.get('subporto', [])),
        '{{POSISI_DIUSULKAN}}': expert.get('posisi_diusulkan', ''),
    }
    
    for placeholder, value in replacements.items():
        replace_placeholder(doc, placeholder, value)
    
    # Handle Pendidikan Formal
    pendidikan_formal = expert.get('pendidikan_formal', [])
    if pendidikan_formal:
        pendidikan_text = '\n'.join([f"• {edu}" for edu in pendidikan_formal])
        replace_placeholder(doc, '{{PENDIDIKAN_FORMAL}}', pendidikan_text)
    
    # Handle Pendidikan Non-Formal
    pendidikan_non_formal = expert.get('pendidikan_non_formal', [])
    if pendidikan_non_formal:
        pelatihan_text = '\n'.join([f"• {training}" for training in pendidikan_non_formal])
        replace_placeholder(doc, '{{PENDIDIKAN_NON_FORMAL}}', pelatihan_text)
    
    # Handle Bahasa
    bahasa = expert.get('penguasaan_bahasa', [])
    if bahasa:
        bahasa_text = '\n'.join([f"• {lang}" for lang in bahasa])
        replace_placeholder(doc, '{{PENGUASAAN_BAHASA}}', bahasa_text)
    
    # Handle Projects - find the project table
    projects = expert.get('projects', [])
    
    # Look for project table markers and fill them
    for table in doc.tables:
        # Check if this is the project table (contains project-related headers)
        first_row_text = ' '.join([cell.text for cell in table.rows[0].cells]).lower()
        
        if 'nama proyek' in first_row_text or 'pemberi kerja' in first_row_text:
            print(f"Found project table with {len(table.rows)} rows")
            
            # Clear existing data rows (keep header)
            # Remove all rows except the first one (header)
            for _ in range(len(table.rows) - 1):
                table._element.remove(table.rows[-1]._element)
            
            # Add project data
            for idx, project in enumerate(projects, 1):
                row_data = [
                    str(idx),
                    project.get('nama_proyek', ''),
                    project.get('pengguna_jasa') or project.get('pemberi_kerja', ''),
                    project.get('lokasi_proyek', ''),
                    str(project.get('tahun', '')),
                    format_currency(project.get('nilai_proyek')),
                    project.get('posisi_penugasan') or project.get('peran', ''),
                    project.get('uraian_tugas', ''),
                    f"{project.get('waktu_mulai', '')} s/d {project.get('waktu_selesai', '')}",
                    project.get('status_kepegawaian', ''),
                    project.get('bersama', ''),
                    project.get('surat_referensi', '')
                ]
                
                # Add row
                row = table.add_row()
                for i, data in enumerate(row_data):
                    if i < len(row.cells):
                        row.cells[i].text = str(data)
            
            print(f"  ✓ Added {len(projects)} projects to table")
            break
    
    # Save document
    doc.save(output_filename)
    print(f"✓ CV saved to: {output_filename}")
    print(f"  - Total Projects: {len(projects)}")
    
    return True

if __name__ == "__main__":
    # Configuration
    expert_id = 730  # Dr. Ir. Budi Santoso
    template_path = "../TEMPLATE_CV_EXPERT.docx"
    output_file = f"CV_Expert_{expert_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    print("=" * 60)
    print("CV GENERATOR FROM TEMPLATE")
    print("=" * 60)
    print()
    
    try:
        success = generate_cv_from_template(expert_id, template_path, output_file)
        if success:
            print()
            print("=" * 60)
            print("SUCCESS!")
            print("=" * 60)
            print(f"\nFile tersimpan: {output_file}")
            print("Silakan buka file untuk melihat hasilnya.")
    except Exception as e:
        print(f"\n❌ ERROR: {e}")
        import traceback
        traceback.print_exc()
