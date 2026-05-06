"""
Generate CV from TEMPLATE_CV_EXPERT.docx following the exact template structure
This script replaces text in specific table cells, not placeholders
"""
import requests
from docx import Document
from datetime import datetime
import os
from copy import deepcopy

BASE_URL = "http://localhost:8000/api/v1"

def format_currency(value):
    """Format currency to Indonesian Rupiah"""
    if not value:
        return "Rp 0"
    return f"Rp {value:,.0f}".replace(",", ".")

def replace_text_in_cell(cell, new_text):
    """Replace all text in a cell while preserving formatting"""
    if cell.text:
        # Clear existing text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = ''
        # Add new text to first paragraph
        if cell.paragraphs:
            cell.paragraphs[0].text = str(new_text)
        else:
            cell.text = str(new_text)

def generate_cv_from_template(expert_id, template_path, output_filename):
    """Generate CV from template for an expert"""
    
    # Check if template exists
    if not os.path.exists(template_path):
        print(f"❌ Template not found: {template_path}")
        return False
    
    # Fetch expert data
    print(f"Fetching data for Expert ID {expert_id}...")
    response = requests.get(f"{BASE_URL}/experts/{expert_id}")
    response.raise_for_status()
    expert = response.json()
    
    print(f"Generating CV for: {expert['nama']}")
    print(f"Loading template: {template_path}")
    
    # Load template
    doc = Document(template_path)
    
    print(f"Template has {len(doc.tables)} tables")
    
    # Table 0: Header Information
    if len(doc.tables) > 0:
        header_table = doc.tables[0]
        print(f"Processing header table (Table 0) with {len(header_table.rows)} rows")
        
        # Row 0: Posisi yang diusulkan
        if len(header_table.rows) > 0 and len(header_table.rows[0].cells) > 3:
            replace_text_in_cell(header_table.rows[0].cells[3], expert.get('posisi_diusulkan', 'Team Leader'))
        
        # Row 2: Nama Personel
        if len(header_table.rows) > 2 and len(header_table.rows[2].cells) > 3:
            replace_text_in_cell(header_table.rows[2].cells[3], expert.get('nama', ''))
        
        # Row 3: Tempat/Tanggal Lahir
        if len(header_table.rows) > 3 and len(header_table.rows[3].cells) > 3:
            ttl = f"{expert.get('tempat_lahir', '')}, {expert.get('tanggal_lahir', '')}"
            replace_text_in_cell(header_table.rows[3].cells[3], ttl)
        
        # Row 4: Pendidikan Formal
        if len(header_table.rows) > 4 and len(header_table.rows[4].cells) > 3:
            pendidikan = expert.get('pendidikan_formal', [])
            if pendidikan:
                pendidikan_text = "\n".join(pendidikan)
                replace_text_in_cell(header_table.rows[4].cells[3], pendidikan_text)
        
        # Row 5: Pendidikan Non Formal
        if len(header_table.rows) > 5 and len(header_table.rows[5].cells) > 3:
            pelatihan = expert.get('pendidikan_non_formal', [])
            if pelatihan:
                pelatihan_text = "\n".join(pelatihan)
                replace_text_in_cell(header_table.rows[5].cells[3], pelatihan_text)
        
        # Row 6: Penguasaan Bahasa
        if len(header_table.rows) > 6 and len(header_table.rows[6].cells) > 3:
            bahasa = expert.get('penguasaan_bahasa', [])
            if bahasa:
                bahasa_text = "\n".join(bahasa)
                replace_text_in_cell(header_table.rows[6].cells[3], bahasa_text)
    
    # Tables 1-3: Project History (maximum 3 projects)
    projects = expert.get('projects', [])[:3]  # Only take first 3 projects
    
    for proj_idx, project in enumerate(projects):
        table_idx = proj_idx + 1  # Tables 1, 2, 3
        
        if table_idx < len(doc.tables):
            proj_table = doc.tables[table_idx]
            print(f"Processing project {proj_idx + 1} (Table {table_idx}) with {len(proj_table.rows)} rows")
            
            # Row 0: Pengalaman Kerja (header - skip)
            
            # Row 1: a. Nama Proyek
            if len(proj_table.rows) > 1 and len(proj_table.rows[1].cells) > 3:
                replace_text_in_cell(proj_table.rows[1].cells[3], project.get('nama_proyek', ''))
            
            # Row 2: b. Lokasi Proyek
            if len(proj_table.rows) > 2 and len(proj_table.rows[2].cells) > 3:
                replace_text_in_cell(proj_table.rows[2].cells[3], project.get('lokasi_proyek', ''))
            
            # Row 3: c. Pengguna Jasa
            if len(proj_table.rows) > 3 and len(proj_table.rows[3].cells) > 3:
                pengguna_jasa = project.get('pengguna_jasa') or project.get('pemberi_kerja', '')
                replace_text_in_cell(proj_table.rows[3].cells[3], pengguna_jasa)
            
            # Row 4: d. Nama Perusahaan
            if len(proj_table.rows) > 4 and len(proj_table.rows[4].cells) > 3:
                perusahaan = project.get('nama_perusahaan_lain') or project.get('bersama', 'PT SUCOFINDO (Persero)')
                replace_text_in_cell(proj_table.rows[4].cells[3], perusahaan)
            
            # Row 5: e. Uraian Tugas
            if len(proj_table.rows) > 5 and len(proj_table.rows[5].cells) > 3:
                replace_text_in_cell(proj_table.rows[5].cells[3], project.get('uraian_tugas', ''))
            
            # Row 6: f. Waktu Pelaksanaan
            if len(proj_table.rows) > 6 and len(proj_table.rows[6].cells) > 3:
                waktu = f"{project.get('waktu_mulai', '')}-{project.get('waktu_selesai', '')}"
                replace_text_in_cell(proj_table.rows[6].cells[3], waktu)
            
            # Row 7: g. Posisi Penugasan
            if len(proj_table.rows) > 7 and len(proj_table.rows[7].cells) > 3:
                posisi = project.get('posisi_penugasan') or project.get('peran', '')
                replace_text_in_cell(proj_table.rows[7].cells[3], posisi)
            
            # Row 8: h. Status Kepegawaian
            if len(proj_table.rows) > 8 and len(proj_table.rows[8].cells) > 3:
                replace_text_in_cell(proj_table.rows[8].cells[3], project.get('status_kepegawaian', 'Tidak Tetap'))
            
            # Row 9: i. Surat Referensi
            if len(proj_table.rows) > 9 and len(proj_table.rows[9].cells) > 3:
                surat_ref = project.get('surat_referensi', '-')
                replace_text_in_cell(proj_table.rows[9].cells[3], surat_ref)
    
    # Last Table: Signature
    last_table_idx = len(doc.tables) - 1
    if last_table_idx >= 0:
        sig_table = doc.tables[last_table_idx]
        print(f"Processing signature table (Table {last_table_idx}) with {len(sig_table.rows)} rows")
        
        # Update tanggal
        if len(sig_table.rows) > 0 and len(sig_table.rows[0].cells) > 1:
            tanggal = datetime.now().strftime('%d %B %Y')
            # Find cell with "Jakarta," and update
            for cell in sig_table.rows[0].cells:
                if 'Jakarta' in cell.text or 'Tanggal' in cell.text:
                    replace_text_in_cell(cell, f"Jakarta, {tanggal}")
                    break
        
        # Update nama
        if len(sig_table.rows) > 2 and len(sig_table.rows[2].cells) > 1:
            replace_text_in_cell(sig_table.rows[2].cells[1], expert.get('nama', ''))
        
        # Update posisi
        if len(sig_table.rows) > 3 and len(sig_table.rows[3].cells) > 1:
            replace_text_in_cell(sig_table.rows[3].cells[1], expert.get('posisi_diusulkan', 'Team Leader'))
    
    # Save document
    doc.save(output_filename)
    print(f"✓ CV saved to: {output_filename}")
    print(f"  - Total Projects in CV: {len(projects)}")
    
    return True

if __name__ == "__main__":
    # Configuration
    expert_id = 730  # Dr. Ir. Budi Santoso
    template_path = "../TEMPLATE_CV_EXPERT.docx"
    output_file = f"CV_{expert_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    print("=" * 70)
    print("CV GENERATOR - MENGGUNAKAN TEMPLATE SUCOFINDO")
    print("=" * 70)
    print()
    
    try:
        success = generate_cv_from_template(expert_id, template_path, output_file)
        if success:
            print()
            print("=" * 70)
            print("✅ SUCCESS!")
            print("=" * 70)
            print(f"\n📄 File tersimpan: {output_file}")
            print("📂 Lokasi: backend/")
            print("\n✨ Silakan buka file untuk melihat hasilnya.")
            print("   Format mengikuti template resmi Sucofindo.")
    except Exception as e:
        print(f"\n❌ ERROR: {e}")
        import traceback
        traceback.print_exc()
