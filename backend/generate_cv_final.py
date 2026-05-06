"""
Generate CV by replacing ALL placeholders in TEMPLATE_CV_EXPERT.docx
Placeholders format: {Placeholder Text}
"""
import requests
from docx import Document
from datetime import datetime
import os
import re

BASE_URL = "http://localhost:8000/api/v1"

def format_currency(value):
    """Format currency to Indonesian Rupiah"""
    if not value:
        return "Rp 0"
    return f"Rp {value:,.0f}".replace(",", ".")

def replace_in_paragraph(paragraph, replacements):
    """Replace placeholders in a paragraph"""
    for placeholder, value in replacements.items():
        if placeholder in paragraph.text:
            # Replace in runs to preserve formatting
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(value))

def replace_in_document(doc, replacements):
    """Replace all placeholders in document (paragraphs and tables)"""
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)

def generate_cv_from_template(expert_id, template_path, output_filename):
    """Generate CV from template for an expert"""
    
    # Check if template exists
    if not os.path.exists(template_path):
        print(f"❌ Template not found: {template_path}")
        return False
    
    # Fetch expert data
    print(f"Fetching data for Expert ID {expert_id}...")
    response = requests.get(f"{BASE_URL}/experts/{expert_id}")
    response.raise_for_status()
    expert = response.json()
    
    print(f"Generating CV for: {expert['nama']}")
    print(f"Loading template: {template_path}")
    
    # Load template
    doc = Document(template_path)
    
    # Prepare replacements for basic info
    replacements = {
        '{Posisi yang Diusulkan}': expert.get('posisi_diusulkan', 'Team Leader'),
        '{Nama Lengkap}': expert.get('nama', ''),
        '{Tempat Lahir}': expert.get('tempat_lahir', ''),
        '{Tanggal Lahir}': expert.get('tanggal_lahir', ''),
        '{Tempat, Tanggal Lahir}': f"{expert.get('tempat_lahir', '')}, {expert.get('tanggal_lahir', '')}",
    }
    
    # Pendidikan Formal
    pendidikan_formal = expert.get('pendidikan_formal', [])
    if pendidikan_formal:
        for i, edu in enumerate(pendidikan_formal, 1):
            replacements[f'{{Pendidikan Formal {i}}}'] = edu
        # Also handle generic placeholder
        replacements['{Pendidikan Formal}'] = '\n'.join(pendidikan_formal)
    
    # Pendidikan Non-Formal
    pendidikan_non_formal = expert.get('pendidikan_non_formal', [])
    if pendidikan_non_formal:
        for i, training in enumerate(pendidikan_non_formal, 1):
            replacements[f'{{Pendidikan Non-Formal {i}}}'] = training
        replacements['{Pendidikan Non-Formal}'] = '\n'.join(pendidikan_non_formal)
    
    # Penguasaan Bahasa
    bahasa = expert.get('penguasaan_bahasa', [])
    if bahasa:
        for i, lang in enumerate(bahasa, 1):
            replacements[f'{{Bahasa {i}}}'] = lang
        replacements['{Penguasaan Bahasa}'] = '\n'.join(bahasa)
    
    # Projects
    projects = expert.get('projects', [])
    
    # Replace for up to 3 projects
    for proj_idx in range(3):
        proj_num = proj_idx + 1
        
        if proj_idx < len(projects):
            project = projects[proj_idx]
            
            # Project-specific replacements
            proj_replacements = {
                f'{{Nama Proyek {proj_num}}}': project.get('nama_proyek', ''),
                f'{{Lokasi Proyek {proj_num}}}': project.get('lokasi_proyek', ''),
                f'{{Nama Klien {proj_num}}}': project.get('pengguna_jasa') or project.get('pemberi_kerja', ''),
                f'{{Nama Perusahaan Tempat Tenaga Ahli Di Hire {proj_num}}}': project.get('nama_perusahaan_lain') or project.get('bersama', 'PT SUCOFINDO (Persero)'),
                f'{{Uraian deskripsi pekerjaan {proj_num}}}': project.get('uraian_tugas', ''),
                f'{{Bulan Awal, Tahun {proj_num}}}': project.get('waktu_mulai', ''),
                f'{{Bulan Akhir, Tahun {proj_num}}}': project.get('waktu_selesai', ''),
                f'{{Durasi {proj_num}}}': f"{project.get('waktu_mulai', '')} - {project.get('waktu_selesai', '')}",
                f'{{Posisi Penugasan {proj_num}}}': project.get('posisi_penugasan') or project.get('peran', ''),
                f'{{Status Kepegawaian {proj_num}}}': project.get('status_kepegawaian', 'Tidak Tetap'),
                f'{{Nomor Surat Referensi {proj_num}}}': project.get('surat_referensi', '-'),
            }
            
            # Also handle without numbers (for generic placeholders)
            if proj_idx == 0:  # First project gets non-numbered placeholders too
                proj_replacements.update({
                    '{Nama Proyek}': project.get('nama_proyek', ''),
                    '{Lokasi Proyek}': project.get('lokasi_proyek', ''),
                    '{Nama Klien}': project.get('pengguna_jasa') or project.get('pemberi_kerja', ''),
                    '{Nama Perusahaan Tempat Tenaga Ahli Di Hire}': project.get('nama_perusahaan_lain') or project.get('bersama', 'PT SUCOFINDO (Persero)'),
                    '{Uraian deskripsi pekerjaan 1}': project.get('uraian_tugas', ''),
                    '{Uraian deskripsi pekerjaan 2}': '',  # Empty if only one description
                    '{Bulan Awal, Tahun}': project.get('waktu_mulai', ''),
                    '{Bulan Akhir, Tahun}': project.get('waktu_selesai', ''),
                    '{Durasi}': f"{project.get('waktu_mulai', '')} - {project.get('waktu_selesai', '')}",
                    '{Posisi Penugasan}': project.get('posisi_penugasan') or project.get('peran', ''),
                    '{Status Kepegawaian}': project.get('status_kepegawaian', 'Tidak Tetap'),
                    '{Nomor Surat Referensi}': project.get('surat_referensi', '-'),
                })
            
            replacements.update(proj_replacements)
        else:
            # Empty placeholders for projects that don't exist
            replacements.update({
                f'{{Nama Proyek {proj_num}}}': '-',
                f'{{Lokasi Proyek {proj_num}}}': '-',
                f'{{Nama Klien {proj_num}}}': '-',
                f'{{Nama Perusahaan Tempat Tenaga Ahli Di Hire {proj_num}}}': '-',
                f'{{Uraian deskripsi pekerjaan {proj_num}}}': '-',
                f'{{Bulan Awal, Tahun {proj_num}}}': '-',
                f'{{Bulan Akhir, Tahun {proj_num}}}': '-',
                f'{{Durasi {proj_num}}}': '-',
                f'{{Posisi Penugasan {proj_num}}}': '-',
                f'{{Status Kepegawaian {proj_num}}}': '-',
                f'{{Nomor Surat Referensi {proj_num}}}': '-',
            })
    
    # Signature section
    tanggal_sekarang = datetime.now().strftime('%d %B %Y')
    replacements.update({
        '{Tanggal}': tanggal_sekarang,
        '{Nama}': expert.get('nama', ''),
        '{Posisi}': expert.get('posisi_diusulkan', 'Team Leader'),
    })
    
    # Perform all replacements
    print(f"Replacing {len(replacements)} placeholders...")
    replace_in_document(doc, replacements)
    
    # Save document
    doc.save(output_filename)
    print(f"✓ CV saved to: {output_filename}")
    print(f"  - Total Projects: {len(projects)}")
    print(f"  - Placeholders replaced: {len(replacements)}")
    
    return True

if __name__ == "__main__":
    # Configuration
    expert_id = 730  # Dr. Ir. Budi Santoso
    template_path = "../TEMPLATE_CV_EXPERT.docx"
    output_file = f"CV_{expert['nama'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx" if 'expert' in locals() else f"CV_Expert_{expert_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    print("=" * 70)
    print("CV GENERATOR - MENGGANTI SEMUA PLACEHOLDER")
    print("=" * 70)
    print()
    
    try:
        # Fetch expert name first for filename
        response = requests.get(f"{BASE_URL}/experts/{expert_id}")
        expert = response.json()
        output_file = f"CV_{expert['nama'].replace(' ', '_').replace('.', '').replace(',', '')}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        success = generate_cv_from_template(expert_id, template_path, output_file)
        if success:
            print()
            print("=" * 70)
            print("✅ SUCCESS!")
            print("=" * 70)
            print(f"\n📄 File tersimpan: {output_file}")
            print("📂 Lokasi: backend/")
            print("\n✨ Semua placeholder sudah diganti dengan data dari database!")
            print("   Silakan buka file untuk verifikasi.")
    except Exception as e:
        print(f"\n❌ ERROR: {e}")
        import traceback
        traceback.print_exc()
