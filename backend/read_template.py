from docx import Document
import sys

try:
    doc = Document('../TEMPLATE_CV_EXPERT.docx')
    
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    print("\n=== PARAGRAPHS ===")
    
    for i, p in enumerate(doc.paragraphs[:50]):
        if p.text.strip():
            print(f"P{i}: {p.text[:200]}")
    
    print("\n=== TABLES ===")
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTable {table_idx}:")
        for row_idx, row in enumerate(table.rows[:10]):
            cells_text = [cell.text[:50] for cell in row.cells]
            print(f"  Row {row_idx}: {' | '.join(cells_text)}")
            
except Exception as e:
    print(f"Error: {e}")
    sys.exit(1)
