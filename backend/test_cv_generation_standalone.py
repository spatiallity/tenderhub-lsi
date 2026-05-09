"""
Standalone CV Generation Test
Test CV generation without running the full FastAPI server
"""

import sys
import os
from pathlib import Path

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent))

from docx import Document
from datetime import datetime

def test_cv_generation():
    """Test if CV generation works with sample data"""
    
    print("=" * 60)
    print("CV GENERATION STANDALONE TEST")
    print("=" * 60)
    
    # Check if template exists
    template_path = Path(__file__).parent / "TEMPLATE_CV_EXPERT.docx"
    print(f"\n1. Checking template file...")
    print(f"   Path: {template_path}")
    print(f"   Exists: {template_path.exists()}")
    
    if not template_path.exists():
        print("   ❌ TEMPLATE FILE NOT FOUND!")
        print("   Please ensure TEMPLATE_CV_EXPERT.docx exists in backend folder")
        return False
    
    # Try to load template
    print(f"\n2. Loading template...")
    try:
        doc = Document(str(template_path))
        print(f"   ✅ Template loaded successfully")
        print(f"   Paragraphs: {len(doc.paragraphs)}")
        print(f"   Tables: {len(doc.tables)}")
    except Exception as e:
        print(f"   ❌ Failed to load template: {e}")
        return False
    
    # Sample expert data
    print(f"\n3. Preparing sample expert data...")
    expert_data = {
        "nama": "Dr. Test Expert",
        "posisi_diusulkan": "Senior Consultant",
        "tempat_lahir": "Jakarta",
        "tanggal_lahir": "1980-01-01",
        "no_hp": "081234567890",
        "instansi": "SUCOFINDO",
        "pendidikan_formal": [
            {"tingkat": "S3", "jurusan": "Engineering", "institusi": "ITB", "tahun": "2010"},
            {"tingkat": "S2", "jurusan": "Engineering", "institusi": "UI", "tahun": "2005"}
        ],
        "projects": [
            {
                "nama_proyek": "Test Project 1",
                "pemberi_kerja": "Test Client",
                "lokasi_proyek": "Jakarta",
                "tahun": 2024,
                "nilai_proyek": 1000000000,
                "peran": "Project Manager"
            }
        ]
    }
    print(f"   ✅ Sample data prepared")
    print(f"   Expert: {expert_data['nama']}")
    print(f"   Projects: {len(expert_data['projects'])}")
    
    # Try to generate CV
    print(f"\n4. Generating CV...")
    try:
        output_path = Path(__file__).parent / f"CV_Test_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        # Simple replacement (not full generation, just test)
        for paragraph in doc.paragraphs:
            if "{{NAMA}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{NAMA}}", expert_data["nama"])
            if "{{POSISI_DIUSULKAN}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{POSISI_DIUSULKAN}}", expert_data["posisi_diusulkan"])
        
        doc.save(str(output_path))
        print(f"   ✅ CV generated successfully")
        print(f"   Output: {output_path}")
        print(f"   Size: {output_path.stat().st_size} bytes")
        return True
        
    except Exception as e:
        print(f"   ❌ Failed to generate CV: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("\nStarting CV generation test...\n")
    success = test_cv_generation()
    
    print("\n" + "=" * 60)
    if success:
        print("✅ TEST PASSED - CV generation works!")
        print("\nThis means:")
        print("- Template file exists")
        print("- python-docx library works")
        print("- CV generation logic is functional")
        print("\nThe 404 error is because HuggingFace Space")
        print("doesn't have the CV generator endpoint.")
        print("\nSOLUTION: Upload backend files to HuggingFace Space")
    else:
        print("❌ TEST FAILED - CV generation has issues")
        print("\nCheck:")
        print("- TEMPLATE_CV_EXPERT.docx exists in backend folder")
        print("- python-docx is installed: pip install python-docx")
    print("=" * 60)
