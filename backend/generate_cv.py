"""
Generate CV from expert data using python-docx
"""
import requests
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

BASE_URL = "http://localhost:8000/api/v1"

def format_currency(value):
    """Format currency to Indonesian Rupiah"""
    if not value:
        return "N/A"
    return f"Rp {value:,.0f}".replace(",", ".")

def generate_cv(expert_id, output_filename):
    """Generate CV document for an expert"""
    
    # Fetch expert data
    print(f"Fetching data for Expert ID {expert_id}...")
    response = requests.get(f"{BASE_URL}/experts/{expert_id}")
    response.raise_for_status()
    expert = response.json()
    
    print(f"Generating CV for: {expert['nama']}")
    
    # Create document
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('CURRICULUM VITAE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Personal Information
    doc.add_heading('DATA PRIBADI', level=1)
    
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    
    personal_data = [
        ('Nama', expert['nama']),
        ('Tempat, Tanggal Lahir', f"{expert.get('tempat_lahir', 'N/A')}, {expert.get('tanggal_lahir', 'N/A')}"),
        ('No. HP', expert.get('no_hp', 'N/A')),
        ('Instansi', expert.get('instansi', 'N/A')),
        ('Keahlian Utama', ', '.join(expert.get('keahlian', []))),
        ('Portfolio', ', '.join(expert.get('subporto', []))),
        ('Posisi Diusulkan', expert.get('posisi_diusulkan', 'N/A'))
    ]
    
    for i, (label, value) in enumerate(personal_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        row.cells[0].width = Inches(2)
    
    doc.add_paragraph()
    
    # Education - Formal
    doc.add_heading('PENDIDIKAN FORMAL', level=1)
    pendidikan_formal = expert.get('pendidikan_formal', [])
    if pendidikan_formal:
        for edu in pendidikan_formal:
            p = doc.add_paragraph(edu, style='List Bullet')
    else:
        doc.add_paragraph('Tidak ada data', style='List Bullet')
    
    doc.add_paragraph()
    
    # Education - Non-Formal
    doc.add_heading('PENDIDIKAN NON-FORMAL / PELATIHAN', level=1)
    pendidikan_non_formal = expert.get('pendidikan_non_formal', [])
    if pendidikan_non_formal:
        for training in pendidikan_non_formal:
            p = doc.add_paragraph(training, style='List Bullet')
    else:
        doc.add_paragraph('Tidak ada data', style='List Bullet')
    
    doc.add_paragraph()
    
    # Language Skills
    doc.add_heading('PENGUASAAN BAHASA', level=1)
    bahasa = expert.get('penguasaan_bahasa', [])
    if bahasa:
        for lang in bahasa:
            doc.add_paragraph(lang, style='List Bullet')
    else:
        doc.add_paragraph('Tidak ada data', style='List Bullet')
    
    doc.add_paragraph()
    
    # Work Experience / Projects
    doc.add_heading('PENGALAMAN KERJA / PROYEK', level=1)
    projects = expert.get('projects', [])
    
    if projects:
        for idx, project in enumerate(projects, 1):
            # Project header
            p = doc.add_paragraph()
            p.add_run(f"{idx}. {project['nama_proyek']}").bold = True
            
            # Project details table
            proj_table = doc.add_table(rows=12, cols=2)
            proj_table.style = 'Light List Accent 1'
            
            project_details = [
                ('Nama Proyek', project.get('nama_proyek', 'N/A')),
                ('Pemberi Kerja / Pengguna Jasa', project.get('pengguna_jasa') or project.get('pemberi_kerja', 'N/A')),
                ('Lokasi Proyek', project.get('lokasi_proyek', 'N/A')),
                ('Tahun', str(project.get('tahun', 'N/A'))),
                ('Nilai Proyek', format_currency(project.get('nilai_proyek'))),
                ('Posisi / Peran', project.get('peran', 'N/A')),
                ('Posisi Penugasan', project.get('posisi_penugasan', 'N/A')),
                ('Status Kepegawaian', project.get('status_kepegawaian', 'N/A')),
                ('Waktu Pelaksanaan', f"{project.get('waktu_mulai', 'N/A')} s/d {project.get('waktu_selesai', 'N/A')}"),
                ('Bersama', project.get('bersama', 'N/A')),
                ('Status Proyek', project.get('status_proyek', 'N/A')),
                ('Surat Referensi', project.get('surat_referensi', 'N/A'))
            ]
            
            for i, (label, value) in enumerate(project_details):
                row = proj_table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = str(value)
                row.cells[0].width = Inches(2.5)
            
            # Uraian Tugas
            doc.add_paragraph()
            doc.add_paragraph('Uraian Tugas:', style='Heading 3')
            uraian = project.get('uraian_tugas', 'Tidak ada uraian tugas')
            doc.add_paragraph(uraian)
            doc.add_paragraph()
    else:
        doc.add_paragraph('Belum ada pengalaman proyek tercatat')
    
    # Reviews (if any)
    reviews = expert.get('reviews', [])
    if reviews:
        doc.add_page_break()
        doc.add_heading('TESTIMONI / REVIEW', level=1)
        for review in reviews:
            p = doc.add_paragraph()
            p.add_run(f"Reviewer: {review['reviewer_nama']}").bold = True
            p.add_run(f" (Rating: {review['rating']}/5)")
            doc.add_paragraph(f'"{review["komentar"]}"', style='Quote')
            doc.add_paragraph()
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer_p = doc.add_paragraph(f"CV ini digenerate otomatis pada {datetime.now().strftime('%d %B %Y, %H:%M WIB')}")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(9)
    footer_p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # Save document
    doc.save(output_filename)
    print(f"✓ CV saved to: {output_filename}")
    print(f"  - Total Projects: {len(projects)}")
    print(f"  - Total Reviews: {len(reviews)}")

if __name__ == "__main__":
    # Generate CV for Dr. Ir. Budi Santoso (ID 730)
    expert_id = 730
    output_file = f"CV_Expert_{expert_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    print("=" * 60)
    print("CV GENERATOR")
    print("=" * 60)
    print()
    
    try:
        generate_cv(expert_id, output_file)
        print()
        print("=" * 60)
        print("SUCCESS!")
        print("=" * 60)
    except Exception as e:
        print(f"ERROR: {e}")
        import traceback
        traceback.print_exc()
